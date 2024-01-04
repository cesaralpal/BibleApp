from docx import Document
from unidecode import unidecode
import json

def analizar_documento(doc_path):
    # Abre el documento de Word
    doc = Document(doc_path)
    map = {}   
    reflexion_tema = False 
    # Imprime el texto de cada párrafo en el documento
    for index, para in enumerate(doc.paragraphs):
        print(f"row {index}: {para.text}")
        if para.text == "":
            continue
        if index == 0:
            map["Semana"] = para.text
        elif index == 10:
            map["Titulo Video"] = para.text
        elif index == 11:
            map["Video Link"] = para.text
        elif index == 12:
            map["Descripcion Video"] = para.text
        elif index == 14:
            map["Titulo Audio"] = para.text
        elif index == 16:         
            map["Descripcion Audio"] = para.text
        elif index == 17:
            map["SoundCloud Link"] = para.text




    for i, table in enumerate(doc.tables):
        for j, row in enumerate(table.rows):
            for k, cell in enumerate(row.cells):
                #print(f"Table {i}, Row {j}, Cell {k}: {cell.text}")

                # Tabla 0, Titulo, fecha, tema, instrucciones
                if i==0 and j==0 and k==0: 
                    #print("Titulo:")
                    #print(cell.text)
                    map["Titulo"] = cell.text
                    #print(f"Table {i}, Row {j}, Cell {k}: {cell.text}")

                elif i==0 and j==1 and k==0:
                    #print("Tema:")
                    #print(cell.text)
                    map["Tema"] = cell.text
                    #print(f"Table {i}, Row {j}, Cell {k}: {cell.text}")

                elif i==0 and j==2 and k==0: 
                    #print("Instrucciones:")
                    #print(cell.text)
                    map["Instrucciones"] = cell.text
                    #print(f"Table {i}, Row {j}, Cell {k}: {cell.text}")

                # Tabla 1, Devocional, Reflexion, Capitulo, *Biografia
                elif i==1 and j==3 and k==0: # Tabla 1, fila 3, celda 0
                    #print("Devocional:")
                    #print(cell.text)
                    map["Devocional"] = cell.text
                    #print(f"Table {i}, Row {j}, Cell {k}: {cell.text}")

                elif i==1 and j==4 and k==0: # Tabla 1, fila 2, celda 0
                    #print("Reflexion:")
                    #print(cell.text)
                    map["Reflexion"] = cell.text   
                    #print(f"Table {i}, Row {j}, Cell {k}: {cell.text}")

                elif i==1 and j==5 and k==0: 
                    #print("Capitulo:")
                    #print(cell.text)
                    map["Capitulo"] = cell.text   
                    #print(f"Table {i}, Row {j}, Cell {k}: {cell.text}")

                elif i==1 and j==6 and k==0: 
                    #print("Lectura:")
                    #print(cell.text)
                    map["Lectura"] = cell.text   
                    #print(f"Table {i}, Row {j}, Cell {k}: {cell.text}")

                elif i==1 and j==8 and k==0: 
                    #print("Biografia:")
                    #print(cell.text)
                    map["Biografia"] = cell.text   
                    #print(f"Table {i}, Row {j}, Cell {k}: {cell.text}")
                elif i==2 and j==1 and k==0:
                    #print("Trivia:")
                    #print(cell.text)
                    trivia= cell.text
                    map["Trivia"] = extract_questions_to_map(trivia)
                    #print(f"Table {i}, Row {j}, Cell {k}: {cell.text}")

    #json_acentos = json.dumps(map, indent = 4)
    #json_sin_acentos = quitar_acentos_en_json(json_acentos)
    #print(json_sin_acentos)
    #print(map)
    return map


def quitar_acentos(texto):
    return unidecode(texto)


def quitar_acentos_en_json(json_data):
    data_dict = json.loads(json_data)
    for key, value in data_dict.items():
        if isinstance(value, str):
            data_dict[key] = quitar_acentos(value)
    return json.dumps(data_dict, indent=4)

import json

def extract_questions_to_map(text):
    lines = text.strip().split('\n')
    questions_map = {}
    current_chapter = None

    # Process each line
    for line in lines:
        # Detect chapter lines by the absence of question marks and presence of numbers
        if '?' not in line and any(char.isdigit() for char in line):
            current_chapter = line.strip()
            questions_map[current_chapter] = []
        elif '?' in line:
            # Append the question to the list within the map under the current chapter
            if current_chapter:
                questions_map[current_chapter].append(line.strip())

    return questions_map

# Reemplaza 'ruta_del_documento.docx' con la ruta de tu propio documento de Word
ruta_del_documento = 'template4.docx'
analizar_documento(ruta_del_documento)